"""
Resume Parser Module

This module provides functionality to parse and extract text from resume files
in different formats (PDF, DOCX).
"""

import os
from typing import Optional, Dict, Any
import logging
import PyPDF2
import docx

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class ResumeParser:
    """Class for parsing resume documents in various formats."""
    
    @staticmethod
    def extract_text_from_pdf(pdf_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text as a string
        """
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:  # Only append if text was extracted
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
        
        return text.strip()

    @staticmethod
    def extract_text_from_docx(docx_path: str) -> str:
        """
        Extract text from a Word document.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text as a string
        """
        text = ""
        try:
            doc = docx.Document(docx_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Only add non-empty paragraphs
                    text += para.text.strip() + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text.strip() + "\n"
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
        
        return text.strip()

    @classmethod
    def parse(cls, file_path: str) -> Optional[Dict[str, Any]]:
        """
        Parse a resume file and extract its text content.
        
        Args:
            file_path: Path to the resume file (PDF or DOCX)
            
        Returns:
            Dictionary containing the extracted text and metadata,
            or None if parsing failed
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
            
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            text = cls.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = cls.extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file format: {file_ext}")
            return None
            
        if not text:
            logger.warning(f"No text extracted from {file_path}")
            return None
            
        # Return the extracted text along with metadata
        return {
            'text': text,
            'file_path': file_path,
            'file_type': file_ext[1:],  # Remove the dot
            'char_count': len(text),
            'word_count': len(text.split())
        }


def parse_resume(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Convenience function to parse a resume file.
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Dictionary with parsed resume data or None if parsing failed
    """
    return ResumeParser.parse(file_path) 